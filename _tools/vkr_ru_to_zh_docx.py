# -*- coding: utf-8 -*-
"""Translate Russian paragraphs in a .docx to zh-CN (Google via deep-translator)."""
from __future__ import annotations

import re
import sys
import time
from pathlib import Path

from docx import Document
from deep_translator import GoogleTranslator

CYR = re.compile(r"[\u0400-\u04FF]")

# Do not translate pure citation lines that are only numbers — still may have RU


def has_cyrillic(s: str) -> bool:
    return bool(CYR.search(s or ""))


def translate_text(translator: GoogleTranslator, text: str, retries: int = 3) -> str:
    t = text.strip()
    if not t or not has_cyrillic(t):
        return text
    for attempt in range(retries):
        try:
            out = translator.translate(t)
            return out if out else text
        except Exception:
            time.sleep(1.5 * (attempt + 1))
    return text


def translate_doc(path_in: Path, path_out: Path) -> None:
    translator = GoogleTranslator(source="ru", target="zh-CN")
    doc = Document(path_in)

    ref_i = None
    for i, p in enumerate(doc.paragraphs):
        if "СПИСОК ИСПОЛЬЗОВАННЫХ" in (p.text or "").upper():
            ref_i = i
            break

    total = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        if ref_i is not None and i > ref_i:
            break
        if not (p.text or "").strip():
            continue
        if not has_cyrillic(p.text):
            continue
        new = translate_text(translator, p.text)
        if new != p.text:
            p.text = new
        if i % 15 == 0:
            print(f"paragraphs {i}/{total}")
        time.sleep(0.25)

    for ti, tbl in enumerate(doc.tables):
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not (p.text or "").strip() or not has_cyrillic(p.text):
                        continue
                    p.text = translate_text(translator, p.text)
                    time.sleep(0.15)
        print(f"table {ti + 1} done")

    doc.save(path_out)
    print(f"Wrote {path_out}")


def main() -> None:
    base = Path(__file__).resolve().parents[1] / "thesis" / "build"
    src = base / "ВКР1_исправлено.docx"
    if len(sys.argv) >= 2:
        src = Path(sys.argv[1])
    out = base / "ВКР1_исправлено_中文.docx"
    if not src.is_file():
        raise SystemExit(f"Нет файла: {src}")
    translate_doc(src, out)


if __name__ == "__main__":
    main()
